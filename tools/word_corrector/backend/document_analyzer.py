import docx
from docx import Document
import mammoth
from spellchecker import SpellChecker
try:
    from hanspell import spell_checker
    HANSPELL_AVAILABLE = True
except ImportError:
    HANSPELL_AVAILABLE = False
    spell_checker = None

try:
    from pykospacing import Spacing
    PYKOSPACING_AVAILABLE = True
except ImportError:
    PYKOSPACING_AVAILABLE = False
    Spacing = None
import re
import time
import logging
from typing import Dict, List, Any
from collections import Counter
from config_manager import config

class DocumentAnalyzer:
    def __init__(self):
        # 영어 맞춤법 검사기
        self.spell_en = SpellChecker()
        # 한국어 띄어쓰기 교정기
        self.spacing = Spacing() if PYKOSPACING_AVAILABLE else None
        # 오류 카운터 및 복구 메커니즘 초기화
        self._korean_error_count = 0
        self._korean_last_reset = time.time()
        self._korean_disabled_logged = False
        
    def analyze_document(self, file_path: str) -> Dict[str, Any]:
        """문서를 종합적으로 분석합니다."""
        
        start_time = time.time()
        max_time = config.get("error_handling.max_analysis_time_seconds", 30)
        
        # 전체 분석이 비활성화된 경우
        if not config.get("analysis_settings.enabled", True):
            return {"message": "Document analysis is disabled", "recommendations": []}
        
        try:
            # 문서 읽기
            doc = Document(file_path)
        except Exception as e:
            if config.get("error_handling.log_errors", True):
                logging.error(f"Failed to read document: {e}")
            raise Exception(f"Failed to read document. Please ensure it's a valid .docx file: {str(e)}")
        
        # HTML 변환 (mammoth 사용)
        try:
            with open(file_path, "rb") as docx_file:
                result = mammoth.convert_to_html(docx_file)
                html_content = result.value
        except Exception as e:
            # HTML 변환 실패 시 계속 진행 (선택적 기능)
            html_content = ""
            if config.get("error_handling.log_errors", True):
                logging.warning(f"HTML conversion failed: {e}")
        
        # 분석 결과 수집
        analysis_result = {
            "filename": file_path.split('/')[-1].split('\\')[-1],
            "syntax_errors": [],
            "spelling_errors": {"korean": [], "english": [], "spacing": []},
            "style_consistency": {},
            "document_structure": {},
            "statistics": {},
            "recommendations": [],
            "analysis_config": {
                "syntax_enabled": config.is_enabled("syntax_errors"),
                "spelling_enabled": config.is_enabled("spelling_errors"),
                "style_enabled": config.is_enabled("style_consistency"),
                "structure_enabled": config.is_enabled("document_structure")
            }
        }
        
        # 각 분석 단계별 실행 (설정에 따라)
        try:
            if config.is_enabled("syntax_errors"):
                analysis_result["syntax_errors"] = self._check_syntax_errors_with_config(doc)
            
            if config.is_enabled("spelling_errors"):
                analysis_result["spelling_errors"] = self._check_spelling_errors_with_config(doc)
            
            if config.is_enabled("style_consistency"):
                analysis_result["style_consistency"] = self._check_style_consistency_with_config(doc)
            
            if config.is_enabled("document_structure"):
                analysis_result["document_structure"] = self._analyze_document_structure_with_config(doc)
            
            # 통계는 항상 생성
            analysis_result["statistics"] = self._get_document_statistics(doc)
            
            # 추천사항 생성
            if config.is_enabled("recommendations"):
                analysis_result["recommendations"] = self._generate_recommendations_with_config(analysis_result)
            
        except Exception as e:
            if config.get("error_handling.continue_on_error", True):
                if config.get("error_handling.log_errors", True):
                    logging.error(f"Analysis error: {e}")
                analysis_result["error"] = str(e)
            else:
                raise e
        
        # 분석 시간 체크
        elapsed_time = time.time() - start_time
        analysis_result["analysis_time"] = elapsed_time
        
        if elapsed_time > max_time:
            logging.warning(f"Analysis took {elapsed_time:.2f}s, exceeding limit of {max_time}s")
        
        return analysis_result
    
    def _check_syntax_errors_with_config(self, doc: Document) -> List[Dict[str, Any]]:
        """설정 기반 구문 오류 검사"""
        errors = []
        max_errors_per_paragraph = config.get_max_errors_per_paragraph("syntax_errors")
        ignore_short = config.should_ignore_short_paragraphs()
        min_length = config.get_min_paragraph_length()
        ignore_patterns = config.get_ignore_patterns("syntax_errors")
        
        # 컴파일된 정규식 패턴들
        compiled_patterns = []
        for pattern in ignore_patterns:
            try:
                compiled_patterns.append(re.compile(pattern))
            except re.error:
                logging.warning(f"Invalid regex pattern: {pattern}")
        
        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            if not text:
                continue
            
            # 짧은 문단 무시 설정 확인
            if ignore_short and len(text) < min_length:
                continue
            
            # 무시 패턴 확인
            should_ignore = False
            for pattern in compiled_patterns:
                if pattern.search(text):
                    should_ignore = True
                    break
            
            if should_ignore:
                continue
            
            paragraph_errors = []
            
            # 각 오류 유형별 검사 (설정에 따라)
            if config.get("syntax_errors.check_punctuation", True):
                paragraph_errors.extend(self._check_punctuation_errors(text, i + 1))
            
            if config.get("syntax_errors.check_brackets", True):
                paragraph_errors.extend(self._check_bracket_errors(text, i + 1))
            
            if config.get("syntax_errors.check_quotes", True):
                paragraph_errors.extend(self._check_quote_errors(text, i + 1))
            
            # 문단당 최대 오류 수 제한
            if len(paragraph_errors) > max_errors_per_paragraph:
                paragraph_errors = paragraph_errors[:max_errors_per_paragraph]
                paragraph_errors.append({
                    "type": "too_many_errors",
                    "paragraph": i + 1,
                    "text": text[:100] + "..." if len(text) > 100 else text,
                    "description": f"이 문단에서 {len(paragraph_errors)}개 이상의 오류가 발견되어 일부만 표시합니다."
                })
            
            errors.extend(paragraph_errors)
        
        return errors
    
    def _check_punctuation_errors(self, text: str, paragraph_num: int) -> List[Dict[str, Any]]:
        """문장 부호 오류 검사"""
        errors = []
        
        # 문장 시작 시 대문자 사용 권장 체크 (설정으로 제어)
        if config.get("syntax_errors.check_punctuation_capitalization", False):
            if re.search(r'[.!?]\s*[a-z가-힣]', text):
                errors.append({
                    "type": "punctuation_capitalization",
                    "paragraph": paragraph_num,
                    "text": text[:100] + "..." if len(text) > 100 else text,
                    "description": "문장 시작 시 대문자 사용 권장"
                })
        
        # 다른 문장 부호 오류들을 여기에 추가할 수 있습니다
        # 예: 연속된 문장부호, 잘못된 문장부호 사용 등
        
        return errors
    
    def _check_bracket_errors(self, text: str, paragraph_num: int) -> List[Dict[str, Any]]:
        """괄호 불일치 검사"""
        errors = []
        open_brackets = text.count('(')
        close_brackets = text.count(')')
        if open_brackets != close_brackets:
            errors.append({
                "type": "brackets",
                "paragraph": paragraph_num,
                "text": text[:100] + "..." if len(text) > 100 else text,
                "description": f"괄호 불일치 (열림: {open_brackets}, 닫힘: {close_brackets})"
            })
        return errors
    
    def _check_quote_errors(self, text: str, paragraph_num: int) -> List[Dict[str, Any]]:
        """따옴표 불일치 검사"""
        errors = []
        quote_count = text.count('"')
        if quote_count % 2 != 0:
            errors.append({
                "type": "quotes",
                "paragraph": paragraph_num,
                "text": text[:100] + "..." if len(text) > 100 else text,
                "description": "따옴표 불일치"
            })
        return errors

    def _check_spelling_errors_with_config(self, doc: Document) -> Dict[str, List[Dict[str, Any]]]:
        """설정 기반 맞춤법 오류 검사"""
        korean_errors = []
        english_errors = []
        spacing_errors = []
        
        ignore_words = set(config.get_ignore_words())
        korean_max = config.get("spelling_errors.korean_max_errors_per_paragraph", 5)
        english_max = config.get("spelling_errors.english_max_errors_per_paragraph", 3)
        spacing_max = config.get("spelling_errors.spacing_max_errors_per_paragraph", 2)
        min_word_length = config.get("spelling_errors.english_min_word_length", 3)
        
        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            if not text:
                continue
            
            paragraph_korean_errors = []
            paragraph_english_errors = []
            paragraph_spacing_errors = []
            
            # 한국어 맞춤법 검사
            if config.get("spelling_errors.korean_enabled", True) and HANSPELL_AVAILABLE:
                # 설정값 로드
                max_text_length = config.get("spelling_errors.korean_max_text_length", 500)
                min_korean_chars = config.get("spelling_errors.korean_min_korean_chars", 5)
                api_delay = config.get("spelling_errors.korean_api_delay", 0.1)
                skip_patterns = config.get("spelling_errors.korean_skip_patterns", [])
                
                # 한국어 글자 수 확인
                korean_char_count = len(re.findall(r'[가-힣]', text))
                korean_check_enabled = korean_char_count >= min_korean_chars
                
                # 제외 패턴 확인
                if korean_check_enabled:
                    for pattern in skip_patterns:
                        try:
                            if re.match(pattern, text):
                                korean_check_enabled = False
                                break
                        except re.error:
                            continue
                
                # 연속 오류 체크 및 복구 메커니즘
                if not hasattr(self, '_korean_error_count'):
                    self._korean_error_count = 0
                if not hasattr(self, '_korean_last_reset'):
                    self._korean_last_reset = time.time()
                
                max_consecutive_errors = config.get("spelling_errors.korean_max_consecutive_errors", 10)
                
                # 5분마다 오류 카운터 리셋 (복구 메커니즘)
                current_time = time.time()
                if current_time - self._korean_last_reset > 300:  # 5분
                    if self._korean_error_count > 0:
                        self._korean_error_count = 0
                        self._korean_last_reset = current_time
                        self._korean_disabled_logged = False  # 로그 플래그 리셋
                        if config.get("error_handling.log_errors", True):
                            logging.info("Korean spell check error counter reset - re-enabling Korean spell check")
                
                # 한국어 검사 실행
                if korean_check_enabled and self._korean_error_count < max_consecutive_errors:
                    try:
                        # 텍스트 길이 제한
                        if len(text) > max_text_length:
                            text_chunks = [text[i:i+max_text_length] for i in range(0, len(text), max_text_length)]
                        else:
                            text_chunks = [text]
                        
                        chunk_success = False
                        for chunk_idx, chunk in enumerate(text_chunks):
                            try:
                                spelled_sent = spell_checker.check(chunk)
                                chunk_success = True
                                
                                # 성공적인 호출 시 오류 카운트 감소 (점진적 복구)
                                if self._korean_error_count > 0:
                                    self._korean_error_count = max(0, self._korean_error_count - 1)
                                
                                # 결과 처리
                                if hasattr(spelled_sent, 'errors') and spelled_sent.errors:
                                    for error in spelled_sent.errors[:korean_max]:
                                        if isinstance(error, (list, tuple)) and len(error) >= 2:
                                            paragraph_korean_errors.append({
                                                "paragraph": i + 1,
                                                "original": str(error[0]),
                                                "suggestion": str(error[1]),
                                                "context": text[:100] + "..." if len(text) > 100 else text
                                            })
                                elif hasattr(spelled_sent, 'checked') and spelled_sent.checked != chunk:
                                    paragraph_korean_errors.append({
                                        "paragraph": i + 1,
                                        "original": chunk,
                                        "suggestion": spelled_sent.checked,
                                        "context": text[:100] + "..." if len(text) > 100 else text
                                    })
                                
                                # API 호출 제한을 위한 대기
                                if chunk_idx < len(text_chunks) - 1:  # 마지막 청크가 아닌 경우만
                                    time.sleep(api_delay)
                                    
                            except Exception as chunk_error:
                                self._korean_error_count += 1
                                if config.get("error_handling.log_errors", True):
                                    logging.debug(f"Korean spell check failed for chunk in paragraph {i+1}: {str(chunk_error)}")
                                break  # 청크 처리 중단
                        
                        # 전체 청크 처리가 실패한 경우에만 경고
                        if not chunk_success:
                            if config.get("error_handling.log_errors", True):
                                logging.debug(f"Korean spell check completely failed for paragraph {i+1}")
                            
                    except Exception as e:
                        self._korean_error_count += 1
                        error_msg = str(e)
                        if config.get("error_handling.log_errors", True):
                            # 'result' 오류는 디버그 레벨로, 다른 오류는 경고 레벨로
                            if "result" in error_msg.lower() or "timeout" in error_msg.lower():
                                logging.debug(f"Korean spell check API issue for paragraph {i+1}: {error_msg}")
                            else:
                                logging.debug(f"Korean spell check failed for paragraph {i+1}: {error_msg}")
                
                elif self._korean_error_count >= max_consecutive_errors:
                    # 한 번만 로그 출력하도록 개선
                    if not hasattr(self, '_korean_disabled_logged') or not self._korean_disabled_logged:
                        if config.get("error_handling.log_errors", True):
                            logging.warning(f"Korean spell check temporarily disabled due to {self._korean_error_count} consecutive errors. Will retry in 5 minutes.")
                        self._korean_disabled_logged = True
            
            # 영어 맞춤법 검사
            if config.get("spelling_errors.english_enabled", True):
                try:
                    english_words = re.findall(r'\b[a-zA-Z]+\b', text)
                    # 길이 필터링 및 무시 단어 제외
                    english_words = [w for w in english_words 
                                   if len(w) >= min_word_length and w.lower() not in ignore_words]
                    
                    misspelled = self.spell_en.unknown(english_words)
                    
                    for word in list(misspelled)[:english_max]:
                        if word.lower() not in ignore_words:
                            candidates = self.spell_en.candidates(word)
                            suggestions = list(candidates)[:3] if candidates else []
                            paragraph_english_errors.append({
                                "paragraph": i + 1,
                                "word": word,
                                "suggestions": suggestions,
                                "context": text[:100] + "..." if len(text) > 100 else text
                            })
                except Exception as e:
                    if config.get("error_handling.log_errors", True):
                        logging.warning(f"English spell check failed for paragraph {i+1}: {e}")
            
            # 한국어 띄어쓰기 검사
            if config.get("spelling_errors.spacing_enabled", True) and PYKOSPACING_AVAILABLE and self.spacing:
                try:
                    corrected = self.spacing(text)
                    if corrected != text:
                        paragraph_spacing_errors.append({
                            "paragraph": i + 1,
                            "original": text,
                            "corrected": corrected,
                            "context": text[:100] + "..." if len(text) > 100 else text
                        })
                        
                        if len(paragraph_spacing_errors) >= spacing_max:
                            break
                except Exception as e:
                    if config.get("error_handling.log_errors", True):
                        logging.warning(f"Spacing check failed for paragraph {i+1}: {e}")
            
            korean_errors.extend(paragraph_korean_errors)
            english_errors.extend(paragraph_english_errors)
            spacing_errors.extend(paragraph_spacing_errors)
        
        return {
            "korean": korean_errors,
            "english": english_errors,
            "spacing": spacing_errors
        }
    
    def _check_style_consistency_with_config(self, doc: Document) -> Dict[str, Any]:
        """설정 기반 스타일 일관성 검사"""
        result = {
            "headings": [],
            "heading_styles": {},
            "font_consistency": {},
            "toc_structure": {"valid": True, "issues": []}
        }
        
        try:
            # 제목 스타일 분석
            if config.get("style_consistency.heading_structure_enabled", True):
                headings = []
                heading_styles = Counter()
                
                for paragraph in doc.paragraphs:
                    if paragraph.style.name.startswith('Heading'):
                        headings.append({
                            "text": paragraph.text,
                            "style": paragraph.style.name,
                            "level": int(paragraph.style.name.split()[-1]) if paragraph.style.name.split()[-1].isdigit() else 1
                        })
                        heading_styles[paragraph.style.name] += 1
                
                result["headings"] = headings
                result["heading_styles"] = dict(heading_styles)
                result["toc_structure"] = self._analyze_toc_structure_with_config(headings)
            
            # 폰트 일관성 검사
            if config.get("style_consistency.font_consistency_enabled", True):
                font_usage = Counter()
                font_size_usage = Counter()
                min_usage = config.get("style_consistency.min_font_usage_threshold", 2)
                
                for paragraph in doc.paragraphs:
                    for run in paragraph.runs:
                        if run.font and run.font.name:
                            font_usage[run.font.name] += 1
                        if run.font and run.font.size:
                            font_size_usage[str(run.font.size)] += 1
                
                # 최소 사용 횟수 이상인 폰트만 포함
                filtered_fonts = {k: v for k, v in font_usage.items() if v >= min_usage}
                filtered_sizes = {k: v for k, v in font_size_usage.items() if v >= min_usage}
                
                result["font_consistency"] = {
                    "fonts_used": filtered_fonts,
                    "font_sizes_used": filtered_sizes,
                    "main_font": font_usage.most_common(1)[0][0] if font_usage else None,
                    "main_font_size": font_size_usage.most_common(1)[0][0] if font_size_usage else None
                }
        
        except Exception as e:
            if config.get("error_handling.log_errors", True):
                logging.error(f"Style consistency check failed: {e}")
            if not config.get("error_handling.continue_on_error", True):
                raise e
        
        return result
    
    def _analyze_toc_structure_with_config(self, headings: List[Dict]) -> Dict[str, Any]:
        """설정 기반 목차 구조 분석"""
        if not headings:
            return {"valid": True, "issues": []}
        
        issues = []
        prev_level = 0
        max_skip = config.get("style_consistency.max_heading_level_skip", 1)
        allow_empty = config.get("style_consistency.allow_empty_headings", False)
        
        for i, heading in enumerate(headings):
            current_level = heading["level"]
            
            # 제목 레벨 건너뛰기 검사
            if current_level > prev_level + max_skip:
                issues.append({
                    "type": "level_skip",
                    "heading": heading["text"],
                    "description": f"제목 레벨이 {prev_level}에서 {current_level}로 너무 많이 건너뛰었습니다."
                })
            
            # 빈 제목 검사
            if not allow_empty and not heading["text"].strip():
                issues.append({
                    "type": "empty_heading",
                    "heading": f"제목 {i+1}",
                    "description": "빈 제목이 발견되었습니다."
                })
            
            prev_level = current_level
        
        return {
            "valid": len(issues) == 0,
            "issues": issues,
            "total_headings": len(headings),
            "max_level": max([h["level"] for h in headings]) if headings else 0
        }
    
    def _analyze_document_structure_with_config(self, doc: Document) -> Dict[str, Any]:
        """설정 기반 문서 구조 분석"""
        total_paragraphs = len(doc.paragraphs)
        
        # 단일 라인 문단 무시 설정
        if config.get("document_structure.ignore_single_line_paragraphs", True):
            non_empty_paragraphs = len([p for p in doc.paragraphs 
                                      if p.text.strip() and len(p.text.strip().split('\n')) > 1])
        else:
            non_empty_paragraphs = len([p for p in doc.paragraphs if p.text.strip()])
        
        empty_paragraphs = total_paragraphs - non_empty_paragraphs
        
        result = {
            "total_paragraphs": total_paragraphs,
            "non_empty_paragraphs": non_empty_paragraphs,
            "empty_paragraphs": empty_paragraphs,
            "table_count": 0,
            "estimated_image_count": 0
        }
        
        # 최소 문단 수 확인
        min_paragraphs = config.get("document_structure.min_paragraphs_for_analysis", 5)
        if total_paragraphs < min_paragraphs:
            result["warning"] = f"문서가 너무 짧습니다 (최소 {min_paragraphs}개 문단 권장)"
        
        # 표 분석
        if config.get("document_structure.analyze_tables", True):
            result["table_count"] = len(doc.tables)
        
        # 이미지 분석
        if config.get("document_structure.analyze_images", True):
            try:
                image_count = 0
                for paragraph in doc.paragraphs:
                    for run in paragraph.runs:
                        if hasattr(run, 'element') and run.element.xpath('.//a:blip'):
                            image_count += 1
                result["estimated_image_count"] = image_count
            except Exception as e:
                if config.get("error_handling.log_errors", True):
                    logging.warning(f"Image analysis failed: {e}")
                result["estimated_image_count"] = 0
        
        return result
    
    def _generate_recommendations_with_config(self, analysis_result: Dict[str, Any]) -> List[str]:
        """설정 기반 추천사항 생성"""
        recommendations = []
        max_recommendations = config.get("recommendations.max_recommendations", 10)
        priority_order = config.get("recommendations.priority_order", 
                                  ["syntax_errors", "spelling_errors", "style_consistency", "document_structure"])
        severity_thresholds = config.get("recommendations.severity_thresholds", 
                                       {"critical": 10, "warning": 5, "info": 1})
        
        # 우선순위에 따라 추천사항 생성
        for section in priority_order:
            if len(recommendations) >= max_recommendations:
                break
                
            if section == "syntax_errors" and analysis_result.get("syntax_errors"):
                error_count = len(analysis_result["syntax_errors"])
                if error_count >= severity_thresholds.get("critical", 10):
                    recommendations.append(f"🚨 심각: 구문 오류 {error_count}개가 발견되었습니다. 즉시 수정이 필요합니다.")
                elif error_count >= severity_thresholds.get("warning", 5):
                    recommendations.append(f"⚠️ 경고: 구문 오류 {error_count}개가 발견되었습니다. 검토를 권장합니다.")
                else:
                    recommendations.append(f"ℹ️ 정보: 구문 오류 {error_count}개가 발견되었습니다.")
            
            elif section == "spelling_errors":
                spelling_data = analysis_result.get("spelling_errors", {})
                total_spelling_errors = (
                    len(spelling_data.get("korean", [])) +
                    len(spelling_data.get("english", [])) +
                    len(spelling_data.get("spacing", []))
                )
                if total_spelling_errors > 0:
                    if total_spelling_errors >= severity_thresholds.get("critical", 10):
                        recommendations.append(f"🚨 심각: 맞춤법 오류 {total_spelling_errors}개가 발견되었습니다.")
                    elif total_spelling_errors >= severity_thresholds.get("warning", 5):
                        recommendations.append(f"⚠️ 경고: 맞춤법 오류 {total_spelling_errors}개가 발견되었습니다.")
                    else:
                        recommendations.append(f"ℹ️ 정보: 맞춤법 오류 {total_spelling_errors}개가 발견되었습니다.")
            
            elif section == "style_consistency":
                style_info = analysis_result.get("style_consistency", {})
                font_info = style_info.get("font_consistency", {})
                max_fonts = config.get("style_consistency.max_different_fonts", 3)
                
                if len(font_info.get("fonts_used", {})) > max_fonts:
                    recommendations.append(f"⚠️ 스타일: 사용된 폰트가 {max_fonts}개를 초과합니다. 일관성을 위해 제한을 권장합니다.")
                
                if not style_info.get("toc_structure", {}).get("valid", True):
                    recommendations.append("⚠️ 구조: 목차 구조에 문제가 있습니다. 제목 레벨을 확인해주세요.")
            
            elif section == "document_structure":
                structure = analysis_result.get("document_structure", {})
                max_empty_ratio = config.get("document_structure.max_empty_paragraph_ratio", 0.4)
                
                if structure.get("total_paragraphs", 0) > 0:
                    empty_ratio = structure.get("empty_paragraphs", 0) / structure["total_paragraphs"]
                    if empty_ratio > max_empty_ratio:
                        recommendations.append(f"⚠️ 구조: 빈 문단 비율이 {empty_ratio*100:.1f}%로 높습니다. 정리를 권장합니다.")
        
        # 추천사항이 없으면 긍정적 메시지
        if not recommendations:
            recommendations.append("✅ 문서가 전반적으로 잘 작성되었습니다!")
        
        return recommendations[:max_recommendations]

    # 기존 메서드들 (호환성을 위해 유지)
    def _get_document_statistics(self, doc: Document) -> Dict[str, Any]:
        """문서 통계를 생성합니다."""
        
        total_text = ""
        for paragraph in doc.paragraphs:
            total_text += paragraph.text + " "
        
        # 단어 수 계산
        korean_chars = len(re.findall(r'[가-힣]', total_text))
        english_words = len(re.findall(r'\b[a-zA-Z]+\b', total_text))
        total_chars = len(total_text.replace(' ', ''))
        
        # 문장 수 계산
        sentences = len(re.findall(r'[.!?]+', total_text))
        
        return {
            "korean_characters": korean_chars,
            "english_words": english_words,
            "total_characters": total_chars,
            "sentences": sentences,
            "estimated_reading_time_minutes": max(1, total_chars // 500)  # 분당 500자 기준
        }