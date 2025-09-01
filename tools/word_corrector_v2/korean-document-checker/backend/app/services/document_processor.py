"""
Document Processor Service

This module provides functionality to extract text, structure, and formatting
information from docx files using python-docx library.
"""

import os
from typing import Dict, List, Optional, Tuple
from pathlib import Path
import logging

from docx import Document
from docx.shared import Inches
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


logger = logging.getLogger(__name__)


class DocumentProcessor:
    """
    Document processor for extracting text, structure, and formatting information
    from docx files.
    """

    def __init__(self):
        """Initialize the document processor."""
        pass

    def extract_text_from_docx(self, file_path: str) -> str:
        """
        Extract plain text content from a docx file.
        
        Args:
            file_path (str): Path to the docx file
            
        Returns:
            str: Extracted text content
            
        Raises:
            FileNotFoundError: If the file doesn't exist
            Exception: If the file cannot be processed
        """
        try:
            if not os.path.exists(file_path):
                raise FileNotFoundError(f"File not found: {file_path}")
            
            doc = Document(file_path)
            text_content = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text.strip())
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_content.append(cell.text.strip())
            
            return "\n".join(text_content)
            
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {str(e)}")
            raise Exception(f"Failed to extract text from document: {str(e)}")

    def extract_structure_info(self, file_path: str) -> Dict:
        """
        Extract document structure information including headings, styles, and layout.
        
        Args:
            file_path (str): Path to the docx file
            
        Returns:
            Dict: Document structure information
        """
        try:
            if not os.path.exists(file_path):
                raise FileNotFoundError(f"File not found: {file_path}")
            
            doc = Document(file_path)
            structure_info = {
                "headings": [],
                "paragraph_styles": [],
                "table_count": 0,
                "total_paragraphs": 0,
                "heading_hierarchy": [],
                "style_usage": {}
            }
            
            # Extract heading information and paragraph styles
            for i, paragraph in enumerate(doc.paragraphs):
                structure_info["total_paragraphs"] += 1
                
                style_name = paragraph.style.name if paragraph.style else "Normal"
                
                # Count style usage
                if style_name in structure_info["style_usage"]:
                    structure_info["style_usage"][style_name] += 1
                else:
                    structure_info["style_usage"][style_name] = 1
                
                # Check if it's a heading
                if style_name.startswith("Heading") or "제목" in style_name:
                    heading_level = self._extract_heading_level(style_name)
                    heading_info = {
                        "text": paragraph.text.strip(),
                        "level": heading_level,
                        "style": style_name,
                        "paragraph_index": i
                    }
                    structure_info["headings"].append(heading_info)
                    structure_info["heading_hierarchy"].append(heading_level)
                
                # Store paragraph style info
                paragraph_style_info = {
                    "index": i,
                    "style": style_name,
                    "text_length": len(paragraph.text),
                    "alignment": self._get_alignment_name(paragraph.alignment)
                }
                structure_info["paragraph_styles"].append(paragraph_style_info)
            
            # Count tables
            structure_info["table_count"] = len(doc.tables)
            
            return structure_info
            
        except Exception as e:
            logger.error(f"Error extracting structure from {file_path}: {str(e)}")
            raise Exception(f"Failed to extract document structure: {str(e)}")

    def get_document_summary(self, file_path: str) -> Dict:
        """
        Get comprehensive document summary including text content and structure.
        
        Args:
            file_path (str): Path to the docx file
            
        Returns:
            Dict: Document summary with content and structure information
        """
        try:
            logger.info(f"Processing document: {file_path}")
            
            # Extract text content
            content = self.extract_text_from_docx(file_path)
            
            # Extract structure information
            structure = self.extract_structure_info(file_path)
            
            # Calculate basic statistics
            word_count = len(content.split()) if content else 0
            character_count = len(content) if content else 0
            paragraph_count = structure.get("total_paragraphs", 0)
            
            summary = {
                "content": content,
                "structure": structure,
                "statistics": {
                    "word_count": word_count,
                    "character_count": character_count,
                    "paragraph_count": paragraph_count,
                    "heading_count": len(structure.get("headings", [])),
                    "table_count": structure.get("table_count", 0)
                },
                "file_info": {
                    "file_path": file_path,
                    "file_name": os.path.basename(file_path),
                    "file_size": os.path.getsize(file_path) if os.path.exists(file_path) else 0
                }
            }
            
            logger.info(f"Document processing completed: {word_count} words, {paragraph_count} paragraphs")
            return summary
            
        except Exception as e:
            logger.error(f"Error getting document summary for {file_path}: {str(e)}")
            raise Exception(f"Failed to get document summary: {str(e)}")

    def _extract_heading_level(self, style_name: str) -> int:
        """
        Extract heading level from style name.
        
        Args:
            style_name (str): Style name
            
        Returns:
            int: Heading level (1-9, default 1)
        """
        try:
            if "Heading" in style_name:
                # Extract number from "Heading 1", "Heading 2", etc.
                parts = style_name.split()
                for part in parts:
                    if part.isdigit():
                        return int(part)
            elif "제목" in style_name:
                # Handle Korean heading styles
                import re
                numbers = re.findall(r'\d+', style_name)
                if numbers:
                    return int(numbers[0])
            
            return 1  # Default heading level
            
        except Exception:
            return 1

    def _get_alignment_name(self, alignment) -> str:
        """
        Get alignment name from alignment enum.
        
        Args:
            alignment: Paragraph alignment enum
            
        Returns:
            str: Alignment name
        """
        try:
            if alignment == WD_PARAGRAPH_ALIGNMENT.LEFT:
                return "left"
            elif alignment == WD_PARAGRAPH_ALIGNMENT.CENTER:
                return "center"
            elif alignment == WD_PARAGRAPH_ALIGNMENT.RIGHT:
                return "right"
            elif alignment == WD_PARAGRAPH_ALIGNMENT.JUSTIFY:
                return "justify"
            else:
                return "left"  # Default
        except Exception:
            return "left"

    def validate_docx_file(self, file_path: str) -> Tuple[bool, Optional[str]]:
        """
        Validate if the file is a valid docx file.
        
        Args:
            file_path (str): Path to the file
            
        Returns:
            Tuple[bool, Optional[str]]: (is_valid, error_message)
        """
        try:
            if not os.path.exists(file_path):
                return False, f"File not found: {file_path}"
            
            # Check file extension
            if not file_path.lower().endswith('.docx'):
                return False, "File is not a .docx file"
            
            # Try to open the document
            doc = Document(file_path)
            
            # Basic validation - check if we can access paragraphs
            _ = len(doc.paragraphs)
            
            return True, None
            
        except Exception as e:
            return False, f"Invalid docx file: {str(e)}"

    def extract_formatting_info(self, file_path: str) -> Dict:
        """
        Extract detailed formatting information from the document.
        
        Args:
            file_path (str): Path to the docx file
            
        Returns:
            Dict: Formatting information
        """
        try:
            if not os.path.exists(file_path):
                raise FileNotFoundError(f"File not found: {file_path}")
            
            doc = Document(file_path)
            formatting_info = {
                "fonts_used": set(),
                "font_sizes": set(),
                "bold_count": 0,
                "italic_count": 0,
                "underline_count": 0,
                "hyperlink_count": 0,
                "list_count": 0
            }
            
            # Analyze paragraph formatting
            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    # Font information
                    if run.font.name:
                        formatting_info["fonts_used"].add(run.font.name)
                    
                    if run.font.size:
                        formatting_info["font_sizes"].add(str(run.font.size))
                    
                    # Style counts
                    if run.bold:
                        formatting_info["bold_count"] += 1
                    if run.italic:
                        formatting_info["italic_count"] += 1
                    if run.underline:
                        formatting_info["underline_count"] += 1
            
            # Convert sets to lists for JSON serialization
            formatting_info["fonts_used"] = list(formatting_info["fonts_used"])
            formatting_info["font_sizes"] = list(formatting_info["font_sizes"])
            
            return formatting_info
            
        except Exception as e:
            logger.error(f"Error extracting formatting info from {file_path}: {str(e)}")
            raise Exception(f"Failed to extract formatting information: {str(e)}")

    def extract_tables_info(self, file_path: str) -> List[Dict]:
        """
        Extract information about tables in the document.
        
        Args:
            file_path (str): Path to the docx file
            
        Returns:
            List[Dict]: List of table information
        """
        try:
            if not os.path.exists(file_path):
                raise FileNotFoundError(f"File not found: {file_path}")
            
            doc = Document(file_path)
            tables_info = []
            
            for i, table in enumerate(doc.tables):
                table_info = {
                    "table_index": i,
                    "row_count": len(table.rows),
                    "column_count": len(table.columns) if table.rows else 0,
                    "cell_count": 0,
                    "has_header": False,
                    "content_preview": []
                }
                
                # Count cells and extract preview content
                for row_idx, row in enumerate(table.rows):
                    row_content = []
                    for cell in row.cells:
                        table_info["cell_count"] += 1
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_content.append(cell_text[:50])  # First 50 chars
                    
                    if row_content and row_idx < 3:  # First 3 rows for preview
                        table_info["content_preview"].append(row_content)
                    
                    # Check if first row might be header (simple heuristic)
                    if row_idx == 0 and row_content:
                        # If first row has shorter text, might be headers
                        avg_length = sum(len(cell) for cell in row_content) / len(row_content)
                        if avg_length < 30:  # Arbitrary threshold
                            table_info["has_header"] = True
                
                tables_info.append(table_info)
            
            return tables_info
            
        except Exception as e:
            logger.error(f"Error extracting tables info from {file_path}: {str(e)}")
            raise Exception(f"Failed to extract tables information: {str(e)}")